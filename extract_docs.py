#!/usr/bin/env python3
"""
Extract text content from DOCX files to understand compliance requirements
"""

import os
from docx import Document

def extract_docx_text(docx_path):
    """Extract all text from a DOCX file"""
    try:
        doc = Document(docx_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        return f"Error extracting {docx_path}: {str(e)}"

def main():
    base_dir = "/Users/maximgavrish/hackathon"
    
    # Extract compliance checklist
    checklist_path = os.path.join(base_dir, "[Issuer] CF Offering Review Checklist[date].docx")
    print("=" * 80)
    print("COMPLIANCE CHECKLIST")
    print("=" * 80)
    if os.path.exists(checklist_path):
        content = extract_docx_text(checklist_path)
        output_path = os.path.join(base_dir, "extracted_checklist.txt")
        with open(output_path, "w") as f:
            f.write(content)
        print(f"✓ Extracted to: {output_path}")
        print(content[:500] + "...\n")
    
    # Extract review prompts
    prompts_path = os.path.join(base_dir, "Form C Review Prompts.docx")
    print("=" * 80)
    print("REVIEW PROMPTS")
    print("=" * 80)
    if os.path.exists(prompts_path):
        content = extract_docx_text(prompts_path)
        output_path = os.path.join(base_dir, "extracted_prompts.txt")
        with open(output_path, "w") as f:
            f.write(content)
        print(f"✓ Extracted to: {output_path}")
        print(content[:500] + "...\n")
    
    # Extract one example from c-forms-issues
    issues_dir = os.path.join(base_dir, "c-forms-issues")
    if os.path.exists(issues_dir):
        print("=" * 80)
        print("EXAMPLE ISSUE FORM (First file)")
        print("=" * 80)
        issue_files = [f for f in os.listdir(issues_dir) if f.endswith('.docx')]
        if issue_files:
            first_issue = os.path.join(issues_dir, issue_files[0])
            content = extract_docx_text(first_issue)
            output_path = os.path.join(base_dir, "extracted_example_issue.txt")
            with open(output_path, "w") as f:
                f.write(content)
            print(f"✓ Extracted {issue_files[0]} to: {output_path}")
            print(content[:500] + "...\n")

if __name__ == "__main__":
    main()

